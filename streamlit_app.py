import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import io

st.title("📝 Aplikasi Rekalkulasi Dokumen Word")
st.write("Upload dokumen Word (.docx) untuk merekalkulasi tabel.")

def recalculate_tables(doc):
    for table in doc.tables:
        if len(table.columns) < 3:
            continue
        
        num_cols = len(table.columns)
        vertical_sums = [0.0] * num_cols
        
        for row in table.rows:
            # Deteksi baris total dengan lebih spesifik
            is_total_row = False
            
            # Cek kolom pertama untuk keyword "Jumlah/Total"
            if len(row.cells) > 0 and "JUMLAH" in row.cells[0].text.upper():
                is_total_row = True
            # Cek jika baris memiliki pola total (kolom 1 dan 2 kosong)
            elif len(row.cells) > 2 and row.cells[0].text.strip() == "" and row.cells[1].text.strip() == "":
                is_total_row = True
            
            if is_total_row:
                continue  # Lewati baris total
            
            # Proses kolom numerik (mulai dari kolom ke-3)
            for col_idx in range(2, num_cols):
                if col_idx >= len(row.cells):
                    continue
                
                cell = row.cells[col_idx]
                value = cell.text.strip().replace('.', '').replace(',', '.')
                
                # Handle tanda strip (-) atau sel kosong
                if value == '-' or value == '':
                    num = 0.0
                elif re.match(r'^-?\d+\.?\d*$', value):
                    num = float(value)
                elif re.match(r'^\(\d+\.?\d*\)$', value):
                    num = -float(value[1:-1])
                else:
                    continue
                
                vertical_sums[col_idx] += num
        
        # Tambahkan baris rekalkulasi
        new_row = table.add_row()
        new_row.cells[0].text = "Rekalkulasi Baru"
        
        for col_idx in range(num_cols):
            if col_idx >= len(new_row.cells):
                break
            
            cell = new_row.cells[col_idx]
            if col_idx >= 2:
                formatted_num = f"{vertical_sums[col_idx]:,.2f}".replace(',', 'temp').replace('.', ',').replace('temp', '.')
                cell.text = formatted_num
            _set_font(cell)

def _set_font(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

uploaded_file = st.file_uploader("Upload File Word (.docx)", type=["docx"])
if uploaded_file:
    try:
        doc = Document(uploaded_file)
        recalculate_tables(doc)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        st.success("Rekalkulasi tabel selesai!")
        st.download_button(
            label="📥 Unduh Hasil",
            data=output,
            file_name="rekalkulasi.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Terjadi kesalahan: {str(e)}")
